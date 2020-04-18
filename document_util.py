# coding=utf-8

from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

from util import hex_to_rgb

SPECIAL_COLOR = RGBColor(0, 0, 215)
NORMAL_COLOR = RGBColor(0, 0, 0)
NONSHIFT = '00000000'


def update_low4(text, update):
    # convert int to binary and update
    bin_text = '{:08b}'.format(text)
    bin_text = bin_text[:-4] + update
    return int(bin_text, 2)


def get_update(text):
    bin_text = '{:08b}'.format(text)
    update = bin_text[-4:]  # 暂时这样写吧
    return update


def set_style(run, mark):
    # run.italic = True
    # run.font.colomr.rgb = SPECIAL_COLOR
    (r, g, b) = NORMAL_COLOR

    color_style = NORMAL_COLOR
    if mark != NONSHIFT:
        g = update_low4(g, mark[:4])
        b = update_low4(b, mark[4:])

        color_style = RGBColor(r, g, b)

    run.font.color.rgb = color_style


def get_special_mark(run):
    rgb_hex = str(run.font.color.rgb)
    (r, g, b) = hex_to_rgb(rgb_hex)
    updateg = get_update(g)
    updateb = get_update(b)
    return updateg + updateb


def read_document(file_path):
    document = Document(file_path)
    content = []
    for p in document.paragraphs:
        content.append(p.text)

    return content


def read_document_style(file_path):
    document = Document(file_path)
    result = dict()
    for p in document.paragraphs:
        if p.text == '':
            continue
        for n in p.runs:
            result[n.text] = get_special_mark(n)

    return result


def write_document(words=[], file_path=None, contents=[]):
    words_dict = {word.word: word for word in words}
    document = Document()

    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    for content in contents:
        p = document.add_paragraph("")
        for word_str in content:
            word = words_dict.get(word_str, None)
            run = p.add_run(word.word.strip())
            # set style
            if word.special_style is True:
                set_style(run, mark=word.mark)
            else:
                set_style(run, mark=NONSHIFT)

    document.save(file_path)


# def read_pdf(file_path):
#     from pdf2image import convert_from_path
#     image = convert_from_path(file_path)
#
#     return image


if __name__ == '__main__':
    # content = read_document("demo.docx")
    # print(content)
    # read_pdf("demo.pdf")
    file_path = 'data/demo_mark.docx'
    read_document_style(file_path)
